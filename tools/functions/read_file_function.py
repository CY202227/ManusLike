import json
import os
import base64
import PyPDF2
from docx import Document
import pandas as pd
import yaml
from openai import OpenAI

class ReadFileFunction:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.client = OpenAI(
            api_key=os.getenv("QwenVl_API_KEY"),
            base_url=os.getenv("QwenVl_BASE_URL"),
        )

    def encode_image(self,file_path:str):
        with open(file_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')

    def read_file(self,file_extension:str):
        with open(self.file_path, "r", encoding="utf-8", errors='ignore') as f:
            content = f.read()
            return {
                    "content": content,
                    "file_type": "text",
                    "extension": file_extension,
                    "size": len(content)
                }
    def read_json_file(self,file_extension:str):
        with open(self.file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            return {
                    "content": json.dumps(data, ensure_ascii=False, indent=2),
                    "file_type": "json",
                    "extension": file_extension
                }
    def read_yaml_file(self,file_extension:str):
        with open(self.file_path, "r", encoding="utf-8") as f:
            content = yaml.safe_load(f)
            return {
                        "content": content,
                        "file_type": "yaml",
                        "extension": file_extension
                    }
    def read_pdf_file(self,file_extension:str):
        with open(self.file_path, "rb") as f:
            try:
                with open(self.file_path, "rb") as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text_content = ""
                    for page in pdf_reader.pages:
                        text_content += page.extract_text()
                return {
                    "content": text_content,
                    "file_type": "pdf",
                    "extension": file_extension,
                    "pages": len(pdf_reader.pages)
                }
            except Exception as e:
                return {"error": f"读取PDF文件失败: {str(e)}"}
    def read_docx_file(self,file_extension:str):
        with open(self.file_path, "rb") as f:
            try:
                doc = Document(f)
                text_content = ""
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                return {
                        "content": text_content,
                        "file_type": "docx",
                        "extension": file_extension,
                        "paragraphs": len(doc.paragraphs)
                    }
            except Exception as e:
                return {"error": f"读取Word文档失败: {str(e)}"}
    def read_xlsx_file(self,file_extension:str):
        try:
            df = pd.read_excel(self.file_path)
            return {
                "content": df.to_dict('records'),
                "file_type": "excel",
                "extension": file_extension,
                "shape": df.shape,
                "columns": df.columns.tolist()
            }
        except Exception as e:
            return {"error": f"读取Excel文件失败: {str(e)}"}
    def read_image_file(self,file_extension:str):
        base64_image = self.encode_image(self.file_path)
        print(f"🖼️ 正在识别图片内容...")
        completion = self.client.chat.completions.create(
            model="qwen-vl-ocr-latest",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": f"data:image/{file_extension};base64,{base64_image}"}}
                    ]
                },
                {
                    "type": "text",
                    "text": "Read all the text in the image."
                }
            ],
            stream=True
        )
        
        # 处理流式响应
        full_response = ""
        for chunk in completion:
            if chunk.choices[0].delta.content:
                content = chunk.choices[0].delta.content
                print(content, end="", flush=True)
                full_response += content
        print()  # 换行
        
        return {
            "content": full_response,
            "file_type": "image",
            "extension": file_extension
        }